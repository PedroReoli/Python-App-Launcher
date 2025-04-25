from docx import Document

# Carregar o documento original
doc = Document("/mnt/data/08_–_FalaAtípica_-_Pedro_Lucas_Reis_-_2025-1_-_V-01_-_Rev[1].docx")

# ----------- APLICAÇÕES MANUAIS DE ALGUMAS MUDANÇAS DIRETAS ----------

# 1. Corrigir as palavras-chave no RESUMO
for para in doc.paragraphs:
    if "Palavras-chave:" in para.text:
        para.text = "Palavras-chave: Gamificação, ScrumBan, Autismo, Estimular Comunicação, Comunicação Alternativa."
        break

# 2. Remover trecho estranho do ABSTRACT e corrigir palavras-chave
for i, para in enumerate(doc.paragraphs):
    if "This Research Project" in para.text:
        doc.paragraphs[i].text = (
            "This research project aims to assist health professionals in stimulating speech in children with speech delays. "
            "The FalaAtipica application uses images, sounds, and gamification elements to encourage verbal communication, "
            "allowing the monitoring of progress and promoting more effective interventions."
        )
    if "Keywords:" in para.text:
        doc.paragraphs[i].text = "Keywords: Gamification, Speech Delay, Autism, Assistive Technology, User Experience"

# 3. Corrigir a lista de siglas em ordem alfabética
siglas_ordenadas = [
    "ABNT - Associação Brasileira de Normas Técnicas",
    "API - Application Programming Interface",
    "CAA - Comunicação Alternativa e Aumentativa",
    "GDD - Game Design Document",
    "IDE - Integrated Development Environment",
    "MVP - Mínimo Produto Viável",
    "PI - Peer Instruction",
    "SCRUMBAN - Método híbrido que combina Scrum e Kanban",
    "SDK - Software Development Kit",
    "UCD - User-Centered Design",
    "UI - User Interface",
    "UX - User Experience"
]
siglas_start = False
for i, para in enumerate(doc.paragraphs):
    if "LISTA DE SIGLAS" in para.text:
        siglas_start = True
        continue
    if siglas_start:
        if para.text.strip() == "":
            break
        doc.paragraphs[i].text = ""
for sigla in siglas_ordenadas:
    doc.add_paragraph(sigla, style="Normal")

# 4. Adicionar texto introdutório no item 1.2
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("1.2 Objetivos específicos"):
        doc.paragraphs[i+1].text = "Para atender ao objetivo geral, é necessário implementar os seguintes objetivos específicos:"
        break

# Salvar documento revisado
output_path_final = "/mnt/data/FalaAtipica_V-02_Revisado.docx"
doc.save(output_path_final)

output_path_final
