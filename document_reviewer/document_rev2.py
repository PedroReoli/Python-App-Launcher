from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Recarregar o documento da versão anterior
doc = Document("/mnt/data/FalaAtipica_V-02_Revisado.docx")

# Função para aplicar recuo de 1,25 cm na primeira linha dos parágrafos
def format_paragraph(paragraph):
    if paragraph.style.name.startswith('Heading'):
        return
    p_format = paragraph.paragraph_format
    p_format.first_line_indent = Pt(28.35)  # equivalente a 1,25 cm
    p_format.left_indent = Pt(0)
    paragraph.style.font.name = 'Times New Roman'
    paragraph.style.font.size = Pt(12)

# Aplicar formatação a todos os parágrafos
for para in doc.paragraphs:
    format_paragraph(para)
    # Corrigir cor da fonte se estiver azul ou colorida
    for run in para.runs:
        run.font.color.rgb = None

# 3. Inserir parágrafo introdutório antes da seção 1 Introdução
for i, para in enumerate(doc.paragraphs):
    if para.text.strip().startswith("1 INTRODUÇÃO"):
        doc.paragraphs[i].insert_paragraph_before("Este trabalho tem como objetivo apresentar o desenvolvimento de um aplicativo assistivo, voltado ao estímulo da fala em crianças com atraso no desenvolvimento da comunicação oral. A seguir, são apresentados os fundamentos teóricos, metodológicos e técnicos que nortearam o projeto.")
        break

# 4. Corrigir título e estrutura da seção “Plano de Desenvolvimento do Jogo”
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "PLANO DE DESENVOLVIMENTO DO JOGO":
        para.text = "4.0 PLANO DE DESENVOLVIMENTO DO JOGO"
        if doc.paragraphs[i+1].text.strip() == "":
            doc.paragraphs[i+1].text = (
                "O desenvolvimento do FalaAtipica foi planejado em etapas estruturadas para garantir a funcionalidade e a eficiência da solução. "
                "O projeto contempla duas versões do aplicativo: uma para as crianças, focada em jogos e estímulos interativos, e outra para os tutores (pais, psicólogos e psiquiatras), "
                "que permitirá o monitoramento e configuração do app infantil. Além disso, haverá um painel web destinado exclusivamente à geração de relatórios detalhados sobre o progresso da criança."
            )
        break

# 5. Adicionar texto introdutório na seção “Escopo do Aplicativo”
for i, para in enumerate(doc.paragraphs):
    if para.text.strip() == "ESCOPO DO APLICATIVO":
        if doc.paragraphs[i+1].text.strip() == "":
            doc.paragraphs[i+1].text = (
                "O escopo do aplicativo FalaAtipica foi dividido em três módulos distintos, cada um com funcionalidades específicas voltadas para facilitar a interação, o monitoramento e a análise do progresso das crianças com atraso de fala."
            )
        break

# Salvar nova versão
output_path_final_v2 = "/mnt/data/FalaAtipica_V-02_Revisado_Completo.docx"
doc.save(output_path_final_v2)

output_path_final_v2
